"""
Read and extract text content from a .docx file.
Usage: python execution/read_docx.py <path_to_file.docx>
"""

import sys
from pathlib import Path


def read_docx(file_path: str) -> str:
    try:
        from docx import Document
    except ImportError:
        print("ERROR: python-docx not installed. Run: pip install python-docx", file=sys.stderr)
        sys.exit(1)

    path = Path(file_path)
    if not path.exists():
        print(f"ERROR: File not found: {file_path}", file=sys.stderr)
        sys.exit(1)

    if path.suffix.lower() != ".docx":
        print(f"ERROR: Expected a .docx file, got: {path.suffix}", file=sys.stderr)
        sys.exit(1)

    doc = Document(path)
    sections = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            sections.append(text)

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                sections.append(row_text)

    if not sections:
        print("WARNING: No text found in document. It may be empty or image-based.", file=sys.stderr)

    return "\n\n".join(sections)


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python execution/read_docx.py <path_to_file.docx>", file=sys.stderr)
        sys.exit(1)

    content = read_docx(sys.argv[1])
    sys.stdout.reconfigure(encoding="utf-8")
    print(content)
